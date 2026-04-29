"""Generic hiring pipeline orchestration — 5 stages + final interview."""

import os
import time

import yaml
from dotenv import load_dotenv

from notion_client import (
    get_candidates,
    get_candidate_data,
    get_code_blocks,
    get_page,
    update_candidate,
    advance_candidate,
    reject_candidate,
    transfer_file_to_notion,
    archive_page,
    patch_page_properties,
    _request,
    _make_rich_text_blocks,
)
from evaluator import evaluate_candidate, evaluate_ranking, evaluate_with_images, fetch_cv_content
from bitcoin_verifier import (
    verify_onchain_transaction,
    verify_lightning_payment,
    match_onchain_by_amount,
    match_lightning_by_amount,
    TARGET_BTC_ADDRESS,
)

load_dotenv(override=True)


INTERVIEW_DB_ID = "33eeddaa-d74c-81e5-ad77-e4861d37dc1c"


def set_email_action(page_id: str, action: str):
    """Set Email Action property to trigger a Notion email automation."""
    update_candidate(page_id, {"Email Action": action})


def add_to_interview_database(candidate: dict, page: dict, stage5_score: int | None = None):
    """Add a candidate who passed Stage 5 to the Interview Candidates database.

    Pulls all stage scores from the original candidate page and creates
    a new entry in the interview database for manual scoring.
    """
    from notion_client import _request

    props = page.get("properties", {})

    # Extract scores from the candidate applications database
    s1 = props.get("AI Score Stage 1", {}).get("number")
    s2 = props.get("Stage 2 Score", {}).get("number")
    s3 = props.get("Stage 3 Score", {}).get("number")
    s4 = props.get("Stage 4 Score", {}).get("number")
    # Use passed-in score (fresh) since page data may be stale
    s5 = stage5_score if stage5_score is not None else props.get("Stage 5 Score", {}).get("number")

    interview_props = {
        "Candidate Name": {"title": [{"text": {"content": candidate.get("full_name", "Unknown")}}]},
        "Email": {"email": candidate.get("email")},
        "LinkedIn": {"url": candidate.get("linkedin")},
        "Interview Status": {"select": {"name": "Pending"}},
        "Source Page": {"url": f"https://www.notion.so/{candidate['page_id'].replace('-', '')}"},
        "Application": {"relation": [{"id": candidate["page_id"]}]},
    }

    # Add scores (skip None values)
    if s1 is not None:
        interview_props["Stage 1 Score"] = {"number": s1}
    if s2 is not None:
        interview_props["Stage 2 Score"] = {"number": s2}
    if s3 is not None:
        interview_props["Stage 3 Score"] = {"number": s3}
    if s4 is not None:
        interview_props["Stage 4 Score"] = {"number": s4}
    if s5 is not None:
        interview_props["Stage 5 Score"] = {"number": s5}

    resp = _request("POST", "/pages", {
        "parent": {"database_id": INTERVIEW_DB_ID},
        "properties": interview_props,
    })
    print(f"  Added to Interview Candidates database: {resp.get('id', 'error')}")
    return resp


def load_config(role: str) -> dict:
    """Load role configuration from YAML file."""
    config_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config")
    config_path = os.path.join(config_dir, f"{role}.yaml")
    with open(config_path) as f:
        return yaml.safe_load(f)


def get_prompts(config: dict) -> list[str]:
    """Read AI prompts from the Notion prompts page."""
    prompts_page_id = config["notion_prompts_page_id"]
    return get_code_blocks(prompts_page_id)


# --- Stage 1: Structured questionnaire + writing test ---

AGE_RANGES_ORDERED = ["Under 25", "25-30", "31-35", "36-40", "41-45", "46-50", "Over 50"]
EXP_RANGES_ORDERED = ["0-2", "3-5", "5-7", "8-10", "10+"]


def _check_hard_filters(candidate: dict, config: dict) -> str | None:
    """Check hard disqualifiers. Returns rejection reason or None if passed."""
    filters = config.get("hard_filters", {})

    # University degree required
    if filters.get("require_degree") and candidate.get("university_degree") != "Yes":
        return "Hard filter: No university degree"

    # Age range check — blank = reject (candidate must provide age)
    max_age = filters.get("max_age_range", "41-45")
    age_range = candidate.get("age_range")
    if not age_range:
        return "Hard filter: Age range not provided"
    max_idx = AGE_RANGES_ORDERED.index(max_age) if max_age in AGE_RANGES_ORDERED else 4
    curr_idx = AGE_RANGES_ORDERED.index(age_range) if age_range in AGE_RANGES_ORDERED else -1
    if curr_idx > max_idx:
        return f"Hard filter: Age range {age_range} exceeds maximum {max_age}"

    # Minimum experience — blank = reject (candidate must provide experience)
    min_exp = filters.get("min_experience", "3-5")
    exp_range = candidate.get("years_sales_range")
    if not exp_range:
        return "Hard filter: Sales experience not provided"
    min_idx = EXP_RANGES_ORDERED.index(min_exp) if min_exp in EXP_RANGES_ORDERED else 1
    curr_idx = EXP_RANGES_ORDERED.index(exp_range) if exp_range in EXP_RANGES_ORDERED else -1
    if curr_idx < min_idx:
        return f"Hard filter: Experience {exp_range} below minimum {min_exp}"

    # Executive sales or closing experience required
    if filters.get("require_executive_or_closing"):
        has_exec = candidate.get("executive_sales") == "Yes"
        has_closing = candidate.get("closed_deals_over_r1m") == "Yes"
        if not has_exec and not has_closing:
            return "Hard filter: No executive sales experience and no deals over R1M"

    return None


# Stage 2+ submission forms each create a NEW row in the Candidate Applications
# database instead of updating the candidate's existing row. The detector and
# merger below recognise those orphan rows by stage and reunite them with the
# original candidate record. Extend this list when a new stage form is added.
# Ordered list of Stage select values used to determine whether a candidate
# has already progressed past a given stage. "Rejected" and other values not
# in this list are treated as unknown (won't block a re-submission).
STAGE_ORDER = [
    "Applied",
    "Stage 1 Review",
    "Stage 2 Task",
    "Stage 3 Task",
    "Stage 4 Task",
    "Stage 5 Task",
    "Final Interview",
    "Hired",
]


def _stage_order_index(stage_name: str | None) -> int:
    """Return the pipeline-order index for a Stage select value, or -1 if
    unknown (e.g. 'Rejected' or None)."""
    if not stage_name:
        return -1
    try:
        return STAGE_ORDER.index(stage_name)
    except ValueError:
        return -1


def _extract_file_text(url: str, filename: str) -> str:
    """Download a file from a (signed) URL and return its plain-text content.

    Supports PDF (via pypdf) and DOCX (via python-docx). Returns "" for any
    other extension or on any failure — the caller should treat extraction as
    best-effort.

    Imports are lazy so that environments without pypdf / python-docx still
    load this module (the merge path falls back to copying the file as-is).
    """
    if not url or not filename:
        return ""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext not in ("pdf", "docx"):
        return ""

    import io
    import httpx
    try:
        with httpx.Client(timeout=60, follow_redirects=True) as client:
            r = client.get(url)
            r.raise_for_status()
            data = r.content
    except Exception as e:
        print(f"  WARN: failed to download {filename} for text extraction: {e}")
        return ""

    try:
        if ext == "pdf":
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(data))
            parts = [(p.extract_text() or "").strip() for p in reader.pages]
            return "\n\n".join(p for p in parts if p).strip()
        if ext == "docx":
            import docx
            d = docx.Document(io.BytesIO(data))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n\n".join(parts).strip()
    except Exception as e:
        print(f"  WARN: failed to extract text from {filename}: {e}")
        return ""
    return ""


STAGE_SUBMISSION_SPECS = [
    {
        "label": "Stage 2",
        "score_prop": "Stage 2 Score",
        "stage_task": "Stage 2 Task",
        "submitted_at_prop": "Stage 2 Submitted At",
        "file_props": [
            "Notion task screenshots",
            "Spreadsheet task screenshots",
            "Presentation task screenshots (each slide)",
            "Upload your task",
        ],
        "text_props": [
            "A.I. email draft (1/3) - prompts",
            "A.I. email draft (2/3) - non edited version",
            "A.I. email draft (3/3) - edited version",
            "Stage 2 Submission",
        ],
    },
    {
        # Stage 3 form uploads a single document (DOCX/PDF) into "Upload your task".
        # On merge, we extract the text and write it into "Stage 3 Submission" so
        # run_stage3 (which only reads the text property) can score it.
        "label": "Stage 3",
        "score_prop": "Stage 3 Score",
        "stage_task": "Stage 3 Task",
        "submitted_at_prop": "Stage 3 Submitted At",
        "file_props": ["Upload your task"],
        "text_props": ["Stage 3 Submission"],
        "extracted_text_target": "Stage 3 Submission",
    },
    {
        # Stage 4 form has 9 separate text fields (one per concept). On merge, we
        # concatenate them into the single "Stage 4 Submission" text property
        # that run_stage4 reads.
        "label": "Stage 4",
        "score_prop": "Stage 4 Score",
        "stage_task": "Stage 4 Task",
        "submitted_at_prop": "Stage 4 Submitted At",
        "file_props": [],
        "text_props": [
            "Bitcoin — What is it and how does it work at a high level?",
            "Lightning Network — What is it and what problem does it solve?",
            "API — What is an API and how is it used in software?",
            "Cloud Services — What are cloud services and why do companies use them?",
            "Web App — What is a web application and how does it differ from a website?",
            "Front End — What does front end mean in software development?",
            "Back End — What does back end mean in software development?",
            "KYC — What is KYC and why is it required?",
            "KYB — What is KYB and how does it differ from KYC?",
        ],
        "concat_text_target": "Stage 4 Submission",
    },
    {
        "label": "Stage 5",
        "score_prop": "Stage 5 Score",
        "stage_task": "Stage 5 Task",
        "submitted_at_prop": "Stage 5 Submitted At",
        "file_props": [
            "Stage 5 BTC Screenshot",
            "On-chain transaction screenshot",
            "Stage 5 Lightning Screenshot",
            "Lightning payment screenshot",
        ],
        "text_props": [
            "Stage 5 BTC Transaction ID",
            "On-chain transaction ID",
            "Stage 5 Lightning Payment Hash",
            "Lightning transaction proof of payment",
            "How many Satoshis did you send on-chain?",
            "How many Satoshis did you send via Lightning?",
        ],
    },
]


def _has_stage1_data(candidate: dict) -> bool:
    """True if the page carries real Stage 1 application data (CV, writing test,
    age, or experience). Orphan submission pages lack all of these."""
    return bool(
        candidate.get("cv_upload")
        or candidate.get("writing_test")
        or candidate.get("age_range")
        or candidate.get("years_sales_range")
    )


def _has_payload_for_spec(page_props: dict, spec: dict) -> bool:
    """True if the page has at least one populated field named in spec."""
    for prop_name in spec["file_props"]:
        if (page_props.get(prop_name) or {}).get("files"):
            return True
    for prop_name in spec["text_props"]:
        rt = (page_props.get(prop_name) or {}).get("rich_text") or []
        if "".join(t.get("plain_text", "") for t in rt).strip():
            return True
    return False


def _detect_stage_submission(candidate: dict, page_props: dict,
                             current_stage: str | None = None) -> dict | None:
    """If this page looks like an orphan submission from a stage form, return
    the matching spec. Otherwise None.

    A page is a stage-N submission orphan iff:
      - it has no Stage 1 data (no CV / writing / age / experience), AND
      - at least one field in the stage's file_props or text_props is populated.

    `current_stage` (the matched original candidate's current Stage value, e.g.
    "Stage 3 Task") disambiguates when multiple specs match the orphan's
    payload. Stage 2 and Stage 3 forms both write to "Upload your task", so
    without this hint we can't tell them apart. With it, we prefer the spec
    whose `stage_task` equals the candidate's current stage.

    Without `current_stage` (e.g. on first detection before the original is
    matched), ties fall back to the LATEST stage in the spec list.
    """
    if _has_stage1_data(candidate):
        return None
    matches = [s for s in STAGE_SUBMISSION_SPECS if _has_payload_for_spec(page_props, s)]
    if not matches:
        return None
    if current_stage:
        for s in matches:
            if s["stage_task"] == current_stage:
                return s
    return matches[-1]


def _normalize_name(s: str | None) -> str:
    """Lowercase, strip, and collapse internal whitespace. Used for name
    comparison so trailing spaces and double-spaces don't break matching."""
    if not s:
        return ""
    return " ".join(s.split()).lower()


def _find_original_candidate(name: str, email: str | None, db_id: str,
                             exclude_page_id: str) -> dict | None:
    """Locate the candidate's existing record in the Candidate Applications
    database, using full name as the primary key and email as a tiebreaker.

    Matching order (per Ricki: full name takes priority; email disambiguates):
      1. Exact normalised full-name match → if exactly one non-archived row,
         that's the match. If multiple, disambiguate by email; if email doesn't
         resolve, prefer the row with an AI Score Stage 1 (the real Stage 1
         record) to avoid matching a prior orphan.
      2. No full-name match → match by email. If exactly one row has this
         email, use it (covers cases where the candidate typed only their
         first name, or a minor spelling variation).
      3. Otherwise None — log [unmatched] for manual review.
    """
    all_rows = get_candidates(db_id)
    candidates = [
        r for r in all_rows
        if r["id"] != exclude_page_id and not r.get("archived") and not r.get("in_trash")
    ]

    target_name = _normalize_name(name)
    target_email = (email or "").strip().lower() or None

    # --- 1. Exact normalised full-name match
    name_matches = []
    for c in candidates:
        cd = get_candidate_data(c)
        if _normalize_name(cd.get("full_name")) == target_name and target_name:
            name_matches.append((c, cd))

    if len(name_matches) == 1:
        return name_matches[0][0]

    if len(name_matches) > 1:
        # Disambiguate by email
        if target_email:
            email_hits = [(c, cd) for (c, cd) in name_matches
                          if (cd.get("email") or "").strip().lower() == target_email]
            if len(email_hits) == 1:
                return email_hits[0][0]
            if len(email_hits) > 1:
                # Prefer the one with a real Stage 1 score
                with_score = [(c, cd) for (c, cd) in email_hits
                              if (c.get("properties", {}).get("AI Score Stage 1", {}) or {}).get("number") is not None]
                if with_score:
                    return with_score[0][0]
                return email_hits[0][0]
        # No email or email didn't help → prefer row with Stage 1 score
        with_score = [(c, cd) for (c, cd) in name_matches
                      if (c.get("properties", {}).get("AI Score Stage 1", {}) or {}).get("number") is not None]
        if with_score:
            return with_score[0][0]
        return None  # ambiguous — don't guess

    # --- 2. No name match → fall back to email
    if target_email:
        email_hits = []
        for c in candidates:
            cd = get_candidate_data(c)
            if (cd.get("email") or "").strip().lower() == target_email:
                email_hits.append((c, cd))
        if len(email_hits) == 1:
            return email_hits[0][0]
        if len(email_hits) > 1:
            with_score = [(c, cd) for (c, cd) in email_hits
                          if (c.get("properties", {}).get("AI Score Stage 1", {}) or {}).get("number") is not None]
            if with_score:
                return with_score[0][0]

    return None


def _is_past_stage(original_props: dict, spec: dict) -> tuple[bool, str | None]:
    """True if the candidate's current Stage is already PAST the stage this
    submission belongs to (game-prevention: can't retroactively boost an
    earlier round's score once you've moved on).

    Returns (is_past, current_stage_name).
    """
    current = (original_props.get("Stage") or {}).get("select") or {}
    current_name = current.get("name")
    current_idx = _stage_order_index(current_name)
    submission_idx = _stage_order_index(spec.get("stage_task"))
    if submission_idx < 0:
        return (False, current_name)
    return (current_idx > submission_idx, current_name)


def _merge_stage_submission(orphan_page: dict, original_page: dict, spec: dict) -> dict:
    """Copy payload fields from the orphan into the original candidate page,
    then archive the orphan.

    Last-wins policy for mid-stage re-submissions: if the candidate is still
    at '<Stage N> Task' and resubmits, the new submission overwrites the old
    and their prior Stage N score is cleared so the per-stage scorer re-runs.

    Game-prevention: if the candidate has already progressed PAST the stage
    this submission belongs to (e.g. they're at Stage 3 Task and resubmit
    Stage 2), the merge is blocked entirely — no fields written, old score
    retained — and the orphan is archived. This stops candidates from
    retroactively replacing a poor earlier-round submission after seeing
    later-round results.

    Files are re-uploaded via Notion's file_uploads API so they persist on the
    target page (signed S3 URLs from the source would otherwise expire).
    """
    orphan_id = orphan_page["id"]
    original_id = original_page["id"]
    orphan_props = orphan_page.get("properties", {})
    original_props = original_page.get("properties", {})
    label = spec["label"]

    patch_props: dict = {}
    overwritten_props: list[str] = []

    # Stamp the stage's Submitted At property with the orphan's creation time
    # (= the moment the candidate hit "submit" on the Notion form). Done early
    # so it's included even if no files/text need patching for some reason.
    submitted_at_prop = spec.get("submitted_at_prop")
    orphan_created_at = orphan_page.get("created_time")
    if submitted_at_prop and orphan_created_at:
        patch_props[submitted_at_prop] = {"date": {"start": orphan_created_at}}

    # --- File properties
    for prop_name in spec["file_props"]:
        src_prop = orphan_props.get(prop_name) or {}
        src_files = src_prop.get("files", [])
        if not src_files:
            continue
        dest_existing = (original_props.get(prop_name) or {}).get("files") or []
        new_entries = []
        for f in src_files:
            ftype = f.get("type")
            if ftype == "file":
                url = f.get("file", {}).get("url")
            elif ftype == "external":
                url = f.get("external", {}).get("url")
            else:
                url = None
            fname = f.get("name") or "file"
            if not url:
                continue
            try:
                new_entries.append(transfer_file_to_notion(url, fname))
            except Exception as e:
                print(f"  WARN: failed to transfer file {fname}: {e}")
        if new_entries:
            patch_props[prop_name] = {"files": new_entries}
            if dest_existing:
                overwritten_props.append(prop_name)

    # --- Text properties
    for prop_name in spec["text_props"]:
        src_prop = orphan_props.get(prop_name) or {}
        text = "".join(t.get("plain_text", "") for t in src_prop.get("rich_text", []))
        if not text.strip():
            continue
        dest_rt = (original_props.get(prop_name) or {}).get("rich_text") or []
        dest_text = "".join(t.get("plain_text", "") for t in dest_rt).strip()
        patch_props[prop_name] = {"rich_text": _make_rich_text_blocks(text)}
        if dest_text:
            overwritten_props.append(prop_name)

    # --- Concatenate text_props into a single target (Stage 4: 9 concept fields
    # → "Stage 4 Submission" so run_stage4 can read it as one blob)
    concat_target = spec.get("concat_text_target")
    if concat_target:
        sections = []
        for prop_name in spec["text_props"]:
            src_prop = orphan_props.get(prop_name) or {}
            t = "".join(rt.get("plain_text", "") for rt in src_prop.get("rich_text", [])).strip()
            if t:
                sections.append(f"## {prop_name}\n{t}")
        if sections:
            combined = "\n\n".join(sections)
            dest_rt = (original_props.get(concat_target) or {}).get("rich_text") or []
            dest_text = "".join(t.get("plain_text", "") for t in dest_rt).strip()
            patch_props[concat_target] = {"rich_text": _make_rich_text_blocks(combined)}
            if dest_text and concat_target not in overwritten_props:
                overwritten_props.append(concat_target)

    # --- Extract text from uploaded file(s) into a target text prop
    # (Stage 3: DOCX/PDF in "Upload your task" → "Stage 3 Submission" so
    # run_stage3 can read it). Best-effort: extraction failure is logged
    # but does not fail the merge.
    extract_target = spec.get("extracted_text_target")
    if extract_target:
        chunks = []
        for prop_name in spec["file_props"]:
            for f in (orphan_props.get(prop_name) or {}).get("files", []):
                ftype = f.get("type")
                if ftype == "file":
                    src_url = (f.get("file") or {}).get("url")
                elif ftype == "external":
                    src_url = (f.get("external") or {}).get("url")
                else:
                    src_url = None
                fname = f.get("name") or ""
                if not src_url:
                    continue
                extracted = _extract_file_text(src_url, fname)
                if extracted:
                    chunks.append(extracted)
        if chunks:
            combined = "\n\n".join(chunks)
            dest_rt = (original_props.get(extract_target) or {}).get("rich_text") or []
            dest_text = "".join(t.get("plain_text", "") for t in dest_rt).strip()
            patch_props[extract_target] = {"rich_text": _make_rich_text_blocks(combined)}
            if dest_text and extract_target not in overwritten_props:
                overwritten_props.append(extract_target)

    # --- Game-prevention: if a re-submission would overwrite existing data
    # AND the candidate has already progressed past this stage, drop the
    # whole merge. The orphan is still archived so the DB stays clean.
    if overwritten_props:
        is_past, current_stage_name = _is_past_stage(original_props, spec)
        if is_past:
            print(f"  [{label} re-submission BLOCKED] candidate is at "
                  f"{current_stage_name!r} (past {spec['stage_task']!r}); "
                  f"retaining prior score and data. Orphan archived.")
            archive_page(orphan_id)
            return {
                "merged": False,
                "blocked_reason": "past_stage",
                "orphan_id": orphan_id,
                "original_id": original_id,
                "stage_label": label,
                "current_stage": current_stage_name,
            }

    # --- Re-submission cleanup: if we overwrote anything (and we're still at
    # or before this stage), clear the prior score + AI writeup and revert
    # Stage so the per-stage scorer picks the candidate up again.
    reset_for_rescore = False
    if overwritten_props:
        score_prop_name = spec.get("score_prop")
        stage_task_value = spec.get("stage_task")
        if score_prop_name:
            patch_props[score_prop_name] = {"number": None}
        # Clear AI writeup (these get overwritten by each stage's evaluator anyway)
        patch_props["AI Reasoning"] = {"rich_text": []}
        patch_props["Strengths"] = {"rich_text": []}
        patch_props["Weaknesses"] = {"rich_text": []}
        if stage_task_value:
            patch_props["Stage"] = {"select": {"name": stage_task_value}}
        reset_for_rescore = True
        print(f"  [{label} re-submission] replaced prior values on {original_id} for: {overwritten_props}")
        print(f"  [{label} re-submission] cleared score + AI writeup; reverted Stage to {stage_task_value!r}")

    if patch_props:
        patch_page_properties(original_id, patch_props)

    # Archive the orphan either way — a re-submission should not leave a dangling row
    archive_page(orphan_id)

    return {
        "merged": bool(patch_props),
        "orphan_id": orphan_id,
        "original_id": original_id,
        "stage_label": label,
        "fields_merged": list(patch_props.keys()),
        "fields_overwritten": overwritten_props,
        "reset_for_rescore": reset_for_rescore,
    }


def _score_single_stage(label: str, config: dict) -> None:
    """Run the per-stage scorer for `label` (e.g. "Stage 2"…"Stage 5").

    The scorers iterate over every candidate currently at the stage, but in
    practice that's usually just the one who just submitted — others have
    either already been scored (idempotent skip) or aren't there yet.
    """
    dispatch = {
        "Stage 2": run_stage2,
        "Stage 3": run_stage3,
        "Stage 4": run_stage4,
        "Stage 5": run_stage5,
    }
    fn = dispatch.get(label)
    if fn is None:
        print(f"  WARN: no scorer registered for {label!r}; skipping inline scoring")
        return
    fn(config)


def process_single_stage1(page_id: str, config: dict) -> dict:
    """Process a single candidate through Stage 1 (hard filters + AI evaluation).

    Called by the webhook server for real-time processing, and by run_stage1()
    for batch processing. Returns a result dict with the outcome.

    Special-case: if the incoming page looks like a Stage 2 task submission
    (from the task submission form, which unfortunately creates a new DB row
    instead of updating the candidate's existing row), merge the submission
    into the original candidate and skip Stage 1 processing.
    """
    thresholds = config["thresholds"]
    stages = config["stages"]
    db_id = config["notion_database_id"]

    # Fetch the single page
    page = get_page(page_id)
    candidate = get_candidate_data(page)
    name = candidate["full_name"]
    email = candidate.get("email")

    # --- Idempotency: if Notion fires the webhook twice or we're asked to
    # re-process an already-archived page, do nothing.
    if page.get("archived") or page.get("is_archived") or page.get("in_trash"):
        return {"page_id": page_id, "name": name, "decision": "Already archived",
                "score": None, "reasoning": "Page is archived; skipping"}

    # --- Stage 2+ submission auto-merge ---
    # Applied BEFORE the idempotency-by-score check and BEFORE hard filters, so
    # Stage 1 disqualifiers (missing CV, missing age, etc.) do not wrongly reject
    # stage submissions that legitimately carry only task payloads.
    spec = _detect_stage_submission(candidate, page.get("properties", {}))
    if spec:
        provisional_label = spec["label"]
        print(f"[{provisional_label} submission provisionally detected] "
              f"page={page_id} name={name!r} email={email!r}")
        original = _find_original_candidate(name, email, db_id, exclude_page_id=page_id)
        if not original:
            msg = (f"{provisional_label} submission for name={name!r} email={email!r} but no matching "
                   f"candidate found — left in place for manual review")
            print(f"  [{provisional_label} unmatched] {msg}")
            return {"page_id": page_id, "name": name, "decision": f"{provisional_label} unmatched",
                    "score": None, "reasoning": msg}
        # Re-run detection with the original's current Stage so Stage 2 vs
        # Stage 3 (which both write to "Upload your task") resolves correctly.
        # Falls back to the provisional spec if no better match exists.
        original_stage = (
            (original.get("properties", {}).get("Stage") or {}).get("select") or {}
        ).get("name")
        refined = _detect_stage_submission(
            candidate, page.get("properties", {}), current_stage=original_stage,
        )
        if refined and refined["label"] != provisional_label:
            print(f"  [disambiguated {provisional_label} → {refined['label']}] "
                  f"candidate is at {original_stage!r}")
        spec = refined or spec
        label = spec["label"]
        merge_result = _merge_stage_submission(page, original, spec)
        if merge_result.get("blocked_reason") == "past_stage":
            reasoning = (f"{label} re-submission blocked: candidate already at "
                         f"{merge_result['current_stage']!r} (past {spec['stage_task']!r}). "
                         f"Prior score retained; orphan archived.")
            return {"page_id": page_id, "name": name, "decision": f"{label} blocked",
                    "score": None, "reasoning": reasoning}
        if merge_result.get("merged"):
            print(f"  Merged into {merge_result['original_id']} "
                  f"(fields: {merge_result['fields_merged']})")
            # Score immediately — GitHub Actions scheduled runs are unreliable
            # (often delayed 1-3 hours). Running inline means candidates see
            # their result within seconds of hitting Submit. If scoring errors,
            # the next cron run still catches them (the scorers are idempotent).
            try:
                _score_single_stage(label, config)
            except Exception as e:
                print(f"  WARN: inline {label} scoring failed ({e}); "
                      f"next cron run will retry")
            reasoning = f"Merged {label} submission into {merge_result['original_id']}; orphan archived"
            if merge_result.get("fields_overwritten"):
                reasoning += (f" (re-submission: overwrote {merge_result['fields_overwritten']}; "
                              f"cleared prior score and reverted to {spec['stage_task']} for re-scoring)")
            return {"page_id": page_id, "name": name, "decision": f"{label} merged",
                    "score": None, "reasoning": reasoning}
        else:
            # Orphan had no usable payload — nothing written, but still archived
            return {"page_id": page_id, "name": name, "decision": f"{label} empty",
                    "score": None,
                    "reasoning": (f"{label} submission had no payload fields populated "
                                  f"for {merge_result['original_id']}; orphan archived")}

    # Stamp Stage 1 Submitted At with the page's creation time (= moment the
    # candidate hit submit on the application form). Done only for real Stage 1
    # applications — stage-submission orphans returned above, so we don't
    # mistakenly tag them.
    stage1_submitted_at = page.get("created_time")
    if stage1_submitted_at:
        try:
            patch_page_properties(page_id, {
                "Stage 1 Submitted At": {"date": {"start": stage1_submitted_at}}
            })
        except Exception as e:
            print(f"  WARN: could not set Stage 1 Submitted At for {page_id}: {e}")

    # Idempotency: skip if already scored
    existing_score = page.get("properties", {}).get("AI Score Stage 1", {}).get("number")
    if existing_score is not None:
        return {"page_id": page_id, "name": name, "decision": "Already processed",
                "score": existing_score, "reasoning": "Skipped — already has Stage 1 score"}

    # Ensure Stage is set to Applied (Notion forms leave it blank)
    current_stage = page.get("properties", {}).get("Stage", {}).get("select")
    if not current_stage:
        update_candidate(page_id, {"Stage": stages["applied"]})

    # Hard filter check BEFORE AI call
    rejection = _check_hard_filters(candidate, config)
    if rejection:
        reject_candidate(page_id, rejection)
        update_candidate(page_id, {
            "AI Score Stage 1": 0,
            "AI Decision Stage 1": "Fail",
            "AI Reasoning": rejection,
            "Red Flags": rejection,
        })
        set_email_action(page_id, "Failed")
        return {"page_id": page_id, "name": name, "decision": "Hard Filter",
                "score": 0, "reasoning": rejection}

    # Load prompts
    prompts = get_prompts(config)
    if not prompts:
        return {"page_id": page_id, "name": name, "decision": "Error",
                "score": None, "reasoning": "No prompts found on the AI Prompts page"}

    stage1_prompt = "\n".join(prompts[:2]) if len(prompts) >= 2 else prompts[0]

    # Fetch CV if available
    cv_urls = candidate.get("cv_upload") or []
    cv_content = fetch_cv_content(cv_urls[0]) if cv_urls else ""

    # Run AI evaluation
    result = evaluate_candidate(candidate, stage1_prompt, cv_content=cv_content)

    score = result.get("score", 0)
    reasoning = result.get("reasoning", "")
    strengths = result.get("strengths", [])
    weaknesses = result.get("weaknesses", [])
    red_flags = result.get("red_flags", [])

    decision = "Pass" if score >= thresholds["stage1_pass"] else "Fail"

    # Update Notion
    update_candidate(page_id, {
        "AI Score Stage 1": score,
        "AI Decision Stage 1": decision,
        "AI Reasoning": reasoning,
        "Strengths": strengths,
        "Weaknesses": weaknesses,
        "Red Flags": red_flags,
    })

    if decision == "Pass":
        advance_candidate(page_id, stages["stage2_task"])
        set_email_action(page_id, "Passed Stage 1")
    else:
        reject_candidate(page_id, reasoning)
        set_email_action(page_id, "Failed")

    return {"page_id": page_id, "name": name, "decision": decision,
            "score": score, "reasoning": reasoning}


def run_stage1(config: dict) -> dict:
    """Screen all 'Applied' candidates through hard filters + AI evaluation."""
    db_id = config["notion_database_id"]
    stages = config["stages"]

    # Get candidates in "Applied" stage AND candidates with no stage set
    # (Notion forms don't auto-set Stage, so new submissions have empty Stage)
    candidates = get_candidates(db_id, stage=stages["applied"])

    from notion_client import _request as _notion_request
    no_stage = _notion_request("POST", f"/databases/{db_id}/query", {
        "filter": {"property": "Stage", "select": {"is_empty": True}},
        "page_size": 100,
    })
    no_stage_pages = no_stage.get("results", [])
    if no_stage_pages:
        # Set them to Applied first, then include them
        for page in no_stage_pages:
            update_candidate(page["id"], {"Stage": stages["applied"]})
        candidates.extend(no_stage_pages)
        print(f"Found {len(no_stage_pages)} new form submission(s) with no stage set — moved to Applied.")

    if not candidates:
        print("No candidates in 'Applied' stage.")
        return {"processed": 0, "passed": 0, "failed": 0, "hard_filtered": 0}

    stats = {"processed": 0, "passed": 0, "failed": 0, "hard_filtered": 0}

    for page in candidates:
        page_id = page["id"]
        name = get_candidate_data(page)["full_name"]
        print(f"\nEvaluating: {name}")

        result = process_single_stage1(page_id, config)
        decision = result["decision"]

        if decision == "Hard Filter":
            stats["hard_filtered"] += 1
            print(f"  HARD FILTER -> Rejected: {result['reasoning']}")
        elif decision == "Pass":
            stats["passed"] += 1
            print(f"  PASS ({result['score']}/100) -> Stage 2")
        elif decision == "Fail":
            stats["failed"] += 1
            print(f"  FAIL ({result['score']}/100) -> Rejected")
        elif decision == "Already processed":
            print(f"  SKIPPED (already scored {result['score']})")
        else:
            print(f"  {decision}: {result['reasoning']}")

        stats["processed"] += 1
        time.sleep(1)

    print(f"\n--- Stage 1 Complete ---")
    print(f"Processed: {stats['processed']}, Hard Filtered: {stats['hard_filtered']}")
    print(f"Passed: {stats['passed']}, Failed: {stats['failed']}")
    return stats


# --- Stage 2: Systems competency (scored /20) ---

def run_stage2(config: dict) -> dict:
    """Evaluate Stage 2 systems competency submissions."""
    db_id = config["notion_database_id"]
    thresholds = config["thresholds"]
    stages = config["stages"]

    candidates = get_candidates(db_id, stage=stages["stage2_task"])
    if not candidates:
        print("No candidates in 'Stage 2 Task' stage.")
        return {"processed": 0, "passed": 0, "failed": 0}

    prompts = get_prompts(config)
    if len(prompts) < 4:
        print("ERROR: Stage 2 prompt not found (expected index 2-3).")
        return {"processed": 0, "passed": 0, "failed": 0}
    stage2_prompt = "\n".join(prompts[2:4])  # Two blocks: scoring + AI detection

    stats = {"processed": 0, "passed": 0, "failed": 0}

    for page in candidates:
        candidate = get_candidate_data(page)
        name = candidate["full_name"]
        print(f"\nEvaluating Stage 2: {name}")

        # Collect submissions per task — each scored independently
        notion_imgs = candidate.get("notion_screenshots", [])
        sheets_imgs = candidate.get("spreadsheet_screenshots", [])
        pres_imgs = candidate.get("presentation_screenshots", [])

        # Build labelled image list so AI knows which screenshots belong to which task
        all_images = []
        image_labels = []
        if notion_imgs:
            for img in notion_imgs:
                all_images.append(img)
                image_labels.append("TASK 1 - Notion Pipeline Screenshot")
        if sheets_imgs:
            for img in sheets_imgs:
                all_images.append(img)
                image_labels.append("TASK 2 - Spreadsheet Screenshot")
        if pres_imgs:
            for img in pres_imgs:
                all_images.append(img)
                image_labels.append("TASK 3 - Presentation Screenshot")

        # Collect text submission for Task 4 (AI email)
        prompts_text = candidate.get("ai_email_prompts", "")
        unedited_text = candidate.get("ai_email_unedited", "")
        edited_text = candidate.get("ai_email_edited", "")

        # Build submission context with clear labelling and missing flags
        text_parts = []
        text_parts.append("=== SUBMISSION STATUS ===")
        text_parts.append(f"Task 1 - Notion Pipeline: {'SUBMITTED ({} screenshot(s))'.format(len(notion_imgs)) if notion_imgs else 'NOT SUBMITTED — score 0 for this task'}")
        text_parts.append(f"Task 2 - Spreadsheet: {'SUBMITTED ({} screenshot(s))'.format(len(sheets_imgs)) if sheets_imgs else 'NOT SUBMITTED — score 0 for this task'}")
        text_parts.append(f"Task 3 - Presentation: {'SUBMITTED ({} screenshot(s))'.format(len(pres_imgs)) if pres_imgs else 'NOT SUBMITTED — score 0 for this task'}")
        text_parts.append(f"Task 4 - AI Email: {'SUBMITTED' if (prompts_text or unedited_text or edited_text) else 'NOT SUBMITTED — score 0 for this task'}")
        text_parts.append("")

        if image_labels:
            text_parts.append("=== IMAGE ORDER ===")
            for i, label in enumerate(image_labels, 1):
                text_parts.append(f"Image {i}: {label}")
            text_parts.append("")

        if prompts_text:
            text_parts.append(f"TASK 4 - AI Email Prompts:\n{prompts_text}")
        if unedited_text:
            text_parts.append(f"TASK 4 - AI Email (unedited AI output):\n{unedited_text}")
        if edited_text:
            text_parts.append(f"TASK 4 - AI Email (human-edited final version):\n{edited_text}")
        text_content = "\n\n".join(text_parts)

        # Fall back to single Stage 2 Submission field if no split fields
        fallback = candidate.get("stage2_submission", "")

        has_images = len(all_images) > 0
        has_text = bool(prompts_text or unedited_text or edited_text or fallback)

        if not has_images and not has_text:
            print(f"  SKIPPED — No Stage 2 submission found for {name}")
            continue

        # Use vision-capable evaluation if screenshots exist
        if has_images:
            print(f"    Sending {len(all_images)} images + text to Claude Vision...")
            print(f"    Tasks submitted: Notion={'Y' if notion_imgs else 'N'}, Sheets={'Y' if sheets_imgs else 'N'}, Pres={'Y' if pres_imgs else 'N'}, Email={'Y' if has_text else 'N'}")
            result = evaluate_with_images(
                prompt=stage2_prompt,
                image_urls=all_images,
                text_content=text_content or fallback,
            )
        else:
            # Text-only evaluation (no screenshots at all)
            submission = text_content or fallback
            result = evaluate_candidate(candidate, stage2_prompt, submission_text=submission)

        score = result.get("score", 0)
        decision = "Pass" if score >= thresholds["stage2_pass"] else "Fail"

        update_candidate(candidate["page_id"], {
            "Stage 2 Score": score,
            "AI Reasoning": result.get("reasoning", ""),
            "Strengths": result.get("strengths", []),
            "Weaknesses": result.get("weaknesses", []),
        })

        if decision == "Pass":
            advance_candidate(candidate["page_id"], stages["stage3_task"])
            set_email_action(candidate["page_id"], "Passed Stage 2")
            stats["passed"] += 1
            print(f"  PASS ({score}/20) -> Stage 3")
        else:
            reject_candidate(candidate["page_id"], result.get("reasoning", ""))
            set_email_action(candidate["page_id"], "Failed")
            stats["failed"] += 1
            print(f"  FAIL ({score}/20) -> Rejected")

        stats["processed"] += 1
        time.sleep(1)

    print(f"\n--- Stage 2 Complete ---")
    print(f"Processed: {stats['processed']}, Passed: {stats['passed']}, Failed: {stats['failed']}")
    return stats


# --- Stage 3: Executive sales simulation (scored /35) ---

def run_stage3(config: dict) -> dict:
    """Evaluate Stage 3 sales simulation submissions."""
    db_id = config["notion_database_id"]
    thresholds = config["thresholds"]
    stages = config["stages"]

    candidates = get_candidates(db_id, stage=stages["stage3_task"])
    if not candidates:
        print("No candidates in 'Stage 3 Task' stage.")
        return {"processed": 0, "passed": 0, "failed": 0}

    prompts = get_prompts(config)
    if len(prompts) < 6:
        print("ERROR: Stage 3 prompt not found (expected index 4-5).")
        return {"processed": 0, "passed": 0, "failed": 0}
    stage3_prompt = "\n".join(prompts[4:6])  # Two blocks: auto-fail checks + scoring

    stats = {"processed": 0, "passed": 0, "failed": 0}

    for page in candidates:
        candidate = get_candidate_data(page)
        name = candidate["full_name"]
        print(f"\nEvaluating Stage 3: {name}")

        submission = candidate.get("stage3_submission", "")
        if not submission:
            cv_urls = candidate.get("cv_upload") or []
            submission = fetch_cv_content(cv_urls[0]) if cv_urls else ""
        if not submission or submission.startswith("["):
            print(f"  SKIPPED — No Stage 3 submission found for {name}")
            continue

        result = evaluate_candidate(candidate, stage3_prompt, submission_text=submission)

        score = result.get("score", 0)
        decision = "Pass" if score >= thresholds["stage3_pass"] else "Fail"

        update_candidate(candidate["page_id"], {
            "Stage 3 Score": score,
            "AI Reasoning": result.get("reasoning", ""),
            "Strengths": result.get("strengths", []),
            "Weaknesses": result.get("weaknesses", []),
        })

        if decision == "Pass":
            advance_candidate(candidate["page_id"], stages["stage4_task"])
            set_email_action(candidate["page_id"], "Passed Stage 3")
            stats["passed"] += 1
            print(f"  PASS ({score}/35) -> Stage 4")
        else:
            reject_candidate(candidate["page_id"], result.get("reasoning", ""))
            set_email_action(candidate["page_id"], "Failed")
            stats["failed"] += 1
            print(f"  FAIL ({score}/35) -> Rejected")

        stats["processed"] += 1
        time.sleep(1)

    print(f"\n--- Stage 3 Complete ---")
    print(f"Processed: {stats['processed']}, Passed: {stats['passed']}, Failed: {stats['failed']}")
    return stats


# --- Stage 4: Technical understanding (scored /25) ---

def run_stage4(config: dict) -> dict:
    """Evaluate Stage 4 technical understanding submissions."""
    db_id = config["notion_database_id"]
    thresholds = config["thresholds"]
    stages = config["stages"]

    candidates = get_candidates(db_id, stage=stages["stage4_task"])
    if not candidates:
        print("No candidates in 'Stage 4 Task' stage.")
        return {"processed": 0, "passed": 0, "failed": 0}

    prompts = get_prompts(config)
    if len(prompts) < 8:
        print("ERROR: Stage 4 prompt not found (expected index 6-7).")
        return {"processed": 0, "passed": 0, "failed": 0}
    stage4_prompt = "\n".join(prompts[6:8])  # Two blocks: reference answers + scoring

    stats = {"processed": 0, "passed": 0, "failed": 0}

    for page in candidates:
        candidate = get_candidate_data(page)
        name = candidate["full_name"]
        print(f"\nEvaluating Stage 4: {name}")

        submission = candidate.get("stage4_submission", "")
        if not submission:
            cv_urls = candidate.get("cv_upload") or []
            submission = fetch_cv_content(cv_urls[0]) if cv_urls else ""
        if not submission or submission.startswith("["):
            print(f"  SKIPPED — No Stage 4 submission found for {name}")
            continue

        result = evaluate_candidate(candidate, stage4_prompt, submission_text=submission)

        score = result.get("score", 0)
        decision = "Pass" if score >= thresholds["stage4_pass"] else "Fail"

        update_candidate(candidate["page_id"], {
            "Stage 4 Score": score,
            "AI Reasoning": result.get("reasoning", ""),
            "Strengths": result.get("strengths", []),
            "Weaknesses": result.get("weaknesses", []),
        })

        if decision == "Pass":
            advance_candidate(candidate["page_id"], stages["stage5_task"])
            set_email_action(candidate["page_id"], "Passed Stage 4")
            stats["passed"] += 1
            print(f"  PASS ({score}/25) -> Stage 5")
        else:
            reject_candidate(candidate["page_id"], result.get("reasoning", ""))
            set_email_action(candidate["page_id"], "Failed")
            stats["failed"] += 1
            print(f"  FAIL ({score}/25) -> Rejected")

        stats["processed"] += 1
        time.sleep(1)

    print(f"\n--- Stage 4 Complete ---")
    print(f"Processed: {stats['processed']}, Passed: {stats['passed']}, Failed: {stats['failed']}")
    return stats


# --- Stage 5: Bitcoin execution (pass/fail) ---


def _build_blockchain_context(
    onchain_data: dict,
    lightning_data: dict,
    txid: str | None,
    payment_hash: str | None,
    onchain_amount_match: dict | None = None,
    lightning_amount_match: dict | None = None,
    claimed_onchain_sats: int | None = None,
    claimed_lightning_sats: int | None = None,
) -> str:
    """Format blockchain verification results as context for the AI prompt."""
    lines = ["=== BLOCKCHAIN VERIFICATION DATA ===", ""]

    # On-chain section
    lines.append("ON-CHAIN (mempool.space):")
    if onchain_data.get("error"):
        lines.append(f"  API Error: {onchain_data['error']}")
    else:
        tx_count = len(onchain_data.get("transactions", []))
        lines.append(f"  Transactions found to target address: {'Yes' if tx_count > 0 else 'No'}")
        lines.append(f"  Transaction count: {tx_count}")
        lines.append(f"  Total received: {onchain_data.get('total_received_sats', 0)} sats")
        if txid:
            match = onchain_data.get("txid_match")
            lines.append(f"  Candidate TXID ({txid[:16]}...): {'VERIFIED - Found on-chain' if match else 'NOT FOUND on-chain'}")
        if claimed_onchain_sats and onchain_amount_match:
            if onchain_amount_match["matched"]:
                best = onchain_amount_match.get("best_match", {})
                lines.append(f"  Candidate claimed {claimed_onchain_sats} sats on-chain: MATCHED (txid: {best.get('txid', 'N/A')[:16]}...)")
            else:
                lines.append(f"  Candidate claimed {claimed_onchain_sats} sats on-chain: NO MATCHING TRANSACTION FOUND")
        for tx in onchain_data.get("transactions", [])[:5]:
            status = "confirmed" if tx["confirmed"] else "unconfirmed"
            lines.append(f"  - txid: {tx['txid'][:16]}... | {tx['amount_sats']} sats | {status}")

    lines.append("")

    # Lightning section
    lines.append("LIGHTNING (Blink API):")
    if lightning_data.get("error"):
        lines.append(f"  API Error: {lightning_data['error']}")
    else:
        pay_count = len(lightning_data.get("recent_payments", []))
        lines.append(f"  Recent incoming payments found: {'Yes' if pay_count > 0 else 'No'}")
        lines.append(f"  Payment count: {pay_count}")
        if payment_hash:
            match = lightning_data.get("hash_match")
            lines.append(f"  Candidate payment hash ({payment_hash[:16]}...): {'VERIFIED - Hash matched' if match else 'NOT FOUND in records'}")
        if claimed_lightning_sats and lightning_amount_match:
            if lightning_amount_match["matched"]:
                best = lightning_amount_match.get("best_match", {})
                lines.append(f"  Candidate claimed {claimed_lightning_sats} sats via Lightning: MATCHED (hash: {(best.get('payment_hash') or 'N/A')[:16]}...)")
            else:
                lines.append(f"  Candidate claimed {claimed_lightning_sats} sats via Lightning: NO MATCHING PAYMENT FOUND")
        for pay in lightning_data.get("recent_payments", [])[:5]:
            lines.append(f"  - {pay['amount_sats']} sats | {pay['created_at']} | hash: {(pay.get('payment_hash') or 'N/A')[:16]}...")

    lines.append("")
    lines.append("VERIFICATION NOTES:")
    lines.append("- Cross-reference the screenshots with the blockchain data above")
    lines.append("- If the candidate provided a transaction ID / payment hash and it was verified, note this as extra credit")
    lines.append("- Amount matching is used to attribute specific transactions to this candidate")
    lines.append(f"- Target BTC address: {TARGET_BTC_ADDRESS}")
    lines.append("- Target Lightning address: useorange@blink.sv")

    return "\n".join(lines)


def _determine_stage5_result(ai_result: dict) -> tuple[str, int]:
    """Extract score and pass/fail from the AI's 6-point rubric evaluation.

    Returns (decision, score) tuple.
    """
    score = ai_result.get("score", 0)
    if not isinstance(score, int):
        try:
            score = int(score)
        except (ValueError, TypeError):
            score = 0

    # Fraud override
    if ai_result.get("fraud_flag"):
        return "Fail", 0

    decision = ai_result.get("decision", "Fail")
    if decision not in ("Pass", "Fail"):
        decision = "Pass" if score >= 4 else "Fail"

    # Enforce threshold regardless of AI decision
    if score < 4:
        decision = "Fail"

    return decision, score


def run_stage5(config: dict) -> dict:
    """Evaluate Stage 5 Bitcoin execution — on-chain + Lightning verification."""
    db_id = config["notion_database_id"]
    stages = config["stages"]

    candidates = get_candidates(db_id, stage=stages["stage5_task"])
    if not candidates:
        print("No candidates in 'Stage 5 Task' stage.")
        return {"processed": 0, "passed": 0, "failed": 0}

    prompts = get_prompts(config)
    if len(prompts) < 10:
        print("ERROR: Stage 5 prompt not found (expected indices 8-9).")
        return {"processed": 0, "passed": 0, "failed": 0}
    stage5_prompt = "\n".join(prompts[8:10])  # Two blocks: scoring rubric + output format

    blink_api_key = os.getenv("BLINK_API_KEY", "")

    stats = {"processed": 0, "passed": 0, "failed": 0}

    for page in candidates:
        candidate = get_candidate_data(page)
        name = candidate["full_name"]
        print(f"\nEvaluating Stage 5: {name}")

        # Gather submissions
        btc_screenshots = candidate.get("stage5_btc_screenshot") or []
        ln_screenshots = candidate.get("stage5_lightning_screenshot") or []
        txid = (candidate.get("stage5_btc_txid") or "").strip() or None
        payment_hash = (candidate.get("stage5_lightning_hash") or "").strip() or None

        # Self-reported amounts for transaction attribution
        # Parse robustly: strip commas, spaces, "sats" suffix, handle BTC notation
        def _parse_sats(raw: str) -> int | None:
            if not raw:
                return None
            # Remove common non-numeric text
            cleaned = raw.lower().replace(",", "").replace(" ", "")
            cleaned = cleaned.replace("sats", "").replace("sat", "").replace("satoshis", "").replace("satoshi", "")
            cleaned = cleaned.strip()
            if not cleaned:
                return None
            try:
                # Handle BTC notation (e.g. "0.00001000")
                val = float(cleaned)
                if val < 1:
                    # Likely BTC, convert to sats
                    return int(round(val * 100_000_000))
                return int(val)
            except ValueError:
                return None

        raw_onchain_sats = (candidate.get("stage5_onchain_sats") or "").strip()
        raw_lightning_sats = (candidate.get("stage5_lightning_sats") or "").strip()
        claimed_onchain_sats = _parse_sats(raw_onchain_sats)
        claimed_lightning_sats = _parse_sats(raw_lightning_sats)

        # Use Notion page last_edited_time as approximate submission timestamp
        submission_timestamp = None
        try:
            from datetime import datetime
            edited = page.get("last_edited_time", "")
            if edited:
                submission_timestamp = int(datetime.fromisoformat(edited.replace("Z", "+00:00")).timestamp())
        except Exception:
            pass

        if not btc_screenshots and not ln_screenshots and not txid and not payment_hash:
            print(f"  SKIPPED — No Stage 5 submissions found for {name}")
            continue

        # Blockchain verification
        print(f"  Verifying on-chain transaction...")
        onchain_data = verify_onchain_transaction(TARGET_BTC_ADDRESS, txid=txid)
        if onchain_data.get("verified"):
            print(f"  On-chain: {len(onchain_data['transactions'])} tx(s) found, {onchain_data['total_received_sats']} sats total")
            if txid and onchain_data.get("txid_match"):
                print(f"  TXID verified on-chain (extra credit)")
        elif onchain_data.get("error"):
            print(f"  On-chain API error: {onchain_data['error']}")
        else:
            print(f"  On-chain: No transactions found to target address")

        print(f"  Verifying Lightning payment...")
        if blink_api_key:
            lightning_data = verify_lightning_payment(blink_api_key, payment_hash=payment_hash)
            if lightning_data.get("verified"):
                print(f"  Lightning: {len(lightning_data['recent_payments'])} incoming payment(s) found")
                if payment_hash and lightning_data.get("hash_match"):
                    print(f"  Payment hash verified (extra credit)")
            elif lightning_data.get("error"):
                print(f"  Lightning API error: {lightning_data['error']}")
            else:
                print(f"  Lightning: No incoming payments found")
        else:
            lightning_data = {"verified": False, "recent_payments": [], "error": "No BLINK_API_KEY configured"}
            print(f"  Lightning: Skipped — no Blink API key")

        # Amount-based transaction attribution
        onchain_amount_match = None
        lightning_amount_match = None
        if claimed_onchain_sats:
            onchain_amount_match = match_onchain_by_amount(onchain_data, claimed_onchain_sats, submission_timestamp)
            if onchain_amount_match["matched"]:
                print(f"  On-chain amount match: {claimed_onchain_sats} sats found")
            else:
                print(f"  On-chain amount match: {claimed_onchain_sats} sats NOT found")
        if claimed_lightning_sats:
            lightning_amount_match = match_lightning_by_amount(lightning_data, claimed_lightning_sats, submission_timestamp)
            if lightning_amount_match["matched"]:
                print(f"  Lightning amount match: {claimed_lightning_sats} sats found")
            else:
                print(f"  Lightning amount match: {claimed_lightning_sats} sats NOT found")

        # Build context and evaluate with AI
        blockchain_context = _build_blockchain_context(
            onchain_data, lightning_data, txid, payment_hash,
            onchain_amount_match=onchain_amount_match,
            lightning_amount_match=lightning_amount_match,
            claimed_onchain_sats=claimed_onchain_sats,
            claimed_lightning_sats=claimed_lightning_sats,
        )
        all_images = (btc_screenshots or []) + (ln_screenshots or [])

        if all_images:
            result = evaluate_with_images(
                prompt=stage5_prompt,
                image_urls=all_images,
                text_content=blockchain_context,
            )
        else:
            # No screenshots — evaluate with blockchain data only
            result = evaluate_candidate(candidate, stage5_prompt, submission_text=blockchain_context)

        # Determine final result from 6-point rubric
        final_result, score = _determine_stage5_result(result)

        # Build detailed reasoning
        reasoning = result.get("reasoning", "")
        item_scores = result.get("item_scores", {})
        if item_scores:
            breakdown = ", ".join(f"{k}={v}" for k, v in item_scores.items())
            reasoning += f"\n[ITEM SCORES] {breakdown}"
        if result.get("fraud_flag"):
            reasoning += "\n[FRAUD FLAG] Screenshots appear fabricated — blockchain data shows no matching transactions."

        update_candidate(candidate["page_id"], {
            "Stage 5 Result": final_result,
            "Stage 5 Score": score,
            "AI Reasoning": reasoning,
        })

        if final_result == "Pass":
            advance_candidate(candidate["page_id"], stages["final_interview"])
            set_email_action(candidate["page_id"], "Passed Stage 5")
            # Only add to interview DB if not already there (duplicate-run protection)
            from notion_client import _request as _nr
            existing = _nr("POST", f"/databases/{INTERVIEW_DB_ID}/query", {
                "filter": {"property": "Source Page", "url": {"equals": f"https://www.notion.so/{candidate['page_id'].replace('-', '')}"}},
                "page_size": 1,
            })
            if not existing.get("results"):
                add_to_interview_database(candidate, page, stage5_score=score)
            else:
                print(f"  Already in Interview DB — skipping duplicate")
            stats["passed"] += 1
            print(f"  PASS ({score}/6) -> Final Interview")
        else:
            reject_candidate(candidate["page_id"], reasoning or "Failed Bitcoin execution test")
            set_email_action(candidate["page_id"], "Failed")
            stats["failed"] += 1
            print(f"  FAIL ({score}/6) -> Rejected")

        stats["processed"] += 1
        time.sleep(1)

    print(f"\n--- Stage 5 Complete ---")
    print(f"Processed: {stats['processed']}, Passed: {stats['passed']}, Failed: {stats['failed']}")
    return stats


# --- Timeout Check ---

def run_timeout_check(config: dict) -> dict:
    """Check active candidates for the application timeout (default 14 days
    from Stage 1 submission).

    Window is read from `config["timeout"]["expiry_days"]` and
    `config["timeout"]["warning_days"]`; the default file ships at
    `expiry_days: 14, warning_days: 12` so candidates get the warning two
    days before they expire. Only applies to candidates in stages 2-5.

    Per-candidate `Extended Deadline` (date property, set manually in Notion)
    overrides the default window: warnings/expiry are measured relative to
    (Extended Deadline + 24h) so an end-of-day override honours the full day.
    """
    from datetime import datetime, timezone, timedelta

    db_id = config["notion_database_id"]
    timeout_cfg = config.get("timeout", {})
    warning_days = timeout_cfg.get("warning_days", 12)
    expiry_days = timeout_cfg.get("expiry_days", 14)
    stages = config["stages"]

    active_stages = [
        stages["stage2_task"],
        stages["stage3_task"],
        stages["stage4_task"],
        stages["stage5_task"],
    ]

    now = datetime.now(timezone.utc)
    stats = {"checked": 0, "warned": 0, "expired": 0, "extended": 0}

    for stage_name in active_stages:
        candidates = get_candidates(db_id, stage=stage_name)
        for page in candidates:
            stats["checked"] += 1
            candidate = get_candidate_data(page)
            name = candidate["full_name"]

            # Read current Email Action to avoid duplicate warnings
            props = page.get("properties", {})
            ea_prop = props.get("Email Action", {}).get("select")
            current_action = ea_prop.get("name") if ea_prop else None

            extended = candidate.get("extended_deadline")
            if extended:
                # Manual override: compute relative to end-of-day on the override date.
                # Date-only strings ("2026-04-24") parse as midnight UTC; +24h gives
                # "end of that day" in UTC so the candidate has the full day.
                try:
                    deadline_at = datetime.fromisoformat(extended.replace("Z", "+00:00"))
                except ValueError:
                    print(f"  SKIP {name}: unparseable Extended Deadline {extended!r}")
                    continue
                if deadline_at.tzinfo is None:
                    deadline_at = deadline_at.replace(tzinfo=timezone.utc)
                deadline_at = deadline_at + timedelta(days=1)
                warning_at = deadline_at - timedelta(days=2)
                stats["extended"] += 1

                if now >= deadline_at:
                    print(f"  TIMEOUT EXPIRED (extended): {name} deadline was {extended}, in {stage_name}")
                    reject_candidate(
                        candidate["page_id"],
                        f"Application timed out: extended deadline {extended} exceeded",
                    )
                    set_email_action(candidate["page_id"], "Timeout Expired")
                    stats["expired"] += 1
                elif now >= warning_at and current_action != "Timeout Warning":
                    print(f"  TIMEOUT WARNING (extended): {name} deadline is {extended}, in {stage_name}")
                    set_email_action(candidate["page_id"], "Timeout Warning")
                    stats["warned"] += 1
                else:
                    print(f"  OK (extended to {extended}): {name} in {stage_name}")
                continue

            # No override — use the configured window from page creation.
            created_str = page.get("created_time", "")
            if not created_str:
                print(f"  SKIP {name}: no created_time")
                continue

            created = datetime.fromisoformat(created_str.replace("Z", "+00:00"))
            elapsed_days = (now - created).total_seconds() / 86400

            if elapsed_days >= expiry_days:
                print(f"  TIMEOUT EXPIRED: {name} ({elapsed_days:.1f} days, in {stage_name})")
                reject_candidate(
                    candidate["page_id"],
                    f"Application timed out: {expiry_days}-day deadline exceeded",
                )
                set_email_action(candidate["page_id"], "Timeout Expired")
                stats["expired"] += 1

            elif elapsed_days >= warning_days and current_action != "Timeout Warning":
                print(f"  TIMEOUT WARNING: {name} ({elapsed_days:.1f} days, in {stage_name})")
                set_email_action(candidate["page_id"], "Timeout Warning")
                stats["warned"] += 1

            else:
                print(f"  OK: {name} ({elapsed_days:.1f} days, in {stage_name})")

        time.sleep(0.35)

    print(f"\n--- Timeout Check Complete ---")
    print(f"Checked: {stats['checked']}, Warned: {stats['warned']}, "
          f"Expired: {stats['expired']}, With extension: {stats['extended']}")
    return stats


# --- Pipeline Health Check ---
#
# Auto-detect stuck submissions and unscored work. Runs as the last step of
# the cron after stages 2-5 + timeout, so any issue created during those
# stages gets caught the same tick. Anything we can't auto-fix lands on the
# `Pipeline Issue` rich-text property — Ricki has a Notion view filtered on
# "Pipeline Issue is not empty" that surfaces them.

PIPELINE_ISSUE_PROP = "Pipeline Issue"
HEALTH_CHECK_GRACE_MINUTES = 5  # how long after submission to wait before flagging unscored


def _set_pipeline_issue(page_id: str, issue: str | None) -> None:
    """Write or clear the Pipeline Issue property on a candidate row."""
    if issue:
        rich = [{"type": "text", "text": {"content": issue[:1900]}}]
        patch_page_properties(page_id, {PIPELINE_ISSUE_PROP: {"rich_text": rich}})
    else:
        patch_page_properties(page_id, {PIPELINE_ISSUE_PROP: {"rich_text": []}})


def _existing_pipeline_issue(page: dict) -> str:
    rt = (page.get("properties", {}).get(PIPELINE_ISSUE_PROP) or {}).get("rich_text") or []
    return "".join(t.get("plain_text", "") for t in rt).strip()


def run_health_check(config: dict) -> dict:
    """Detect stuck submissions; auto-fix where safe, flag the rest.

    Four scans:
      A. Orphan rows (no CV / no Stage 1 score / no age) carrying submission
         payload → re-run process_single_stage1 to merge them.
      B. Rejected candidates with a Stage N submission timestamp but no
         Stage N score → flag for manual review (we don't auto-revive
         a rejected candidate).
      C. Candidates at "Stage N Task" with a Stage N submission >5 min old
         and no Stage N score → call the per-stage scorer once (it's
         idempotent). Re-check; if still unscored, flag.
      D. Candidates at "Stage N Task" (N≥3) with no Stage N-1 score → flag
         broken progression unless their AI Reasoning records a manual
         override or revival.

    Anything currently flagged that no longer matches a condition gets its
    `Pipeline Issue` cleared.
    """
    db_id = config["notion_database_id"]

    print("\n=== Pipeline Health Check ===")
    all_pages = get_candidates(db_id)

    issues: dict[str, str] = {}        # page_id -> issue text
    orphan_ids: list[str] = []
    needs_rescore: dict[int, list[str]] = {}  # stage_n -> [page_ids]

    now_utc = _utcnow_for_health()

    for page in all_pages:
        pr = page.get("properties", {})
        page_id = page["id"]
        if page.get("archived") or page.get("in_trash"):
            continue
        candidate = get_candidate_data(page)
        name = candidate.get("full_name") or "(no name)"
        cv_count = len(candidate.get("cv_upload") or [])
        stage1_score = (pr.get("AI Score Stage 1") or {}).get("number")
        age_select = (pr.get("How old are you?") or {}).get("select")
        stage_select = (pr.get("Stage") or {}).get("select") or {}
        stage = stage_select.get("name")

        # A. Orphan
        is_orphan = (cv_count == 0 and stage1_score is None and not age_select)
        if is_orphan:
            # Check it has a submission-shape payload — otherwise skip (might
            # just be a manually-created blank row)
            if _has_any_stage_payload(pr):
                orphan_ids.append(page_id)
                continue  # we'll process orphans as a batch below

        # B. Rejected with submission but no score
        if stage == "Rejected":
            for n in (2, 3, 4, 5):
                submitted = ((pr.get(f"Stage {n} Submitted At") or {}).get("date") or {}).get("start")
                score = (pr.get(f"Stage {n} Score") or {}).get("number")
                if submitted and score is None:
                    issues[page_id] = (f"Rejected candidate has unscored Stage {n} submission "
                                       f"({submitted}). Manual review: revive or confirm rejection.")
                    break
            continue

        # C. Stage N Task with stale unscored submission
        if stage and "Task" in stage:
            try:
                n = int(stage.split(" ")[1])
            except (IndexError, ValueError):
                n = None
            if n is not None and 2 <= n <= 5:
                submitted = ((pr.get(f"Stage {n} Submitted At") or {}).get("date") or {}).get("start")
                score = (pr.get(f"Stage {n} Score") or {}).get("number")
                if submitted and score is None:
                    age_min = _minutes_since(submitted, now_utc)
                    if age_min is not None and age_min > HEALTH_CHECK_GRACE_MINUTES:
                        needs_rescore.setdefault(n, []).append(page_id)
                        # We'll set/clear the issue after re-scoring below

                # D. Broken progression (only flag once we've confirmed C didn't trigger)
                if n >= 3 and (pr.get(f"Stage {n-1} Score") or {}).get("number") is None:
                    reasoning = "".join(
                        t.get("plain_text", "") for t in (pr.get("AI Reasoning") or {}).get("rich_text", [])
                    )
                    if "MANUAL OVERRIDE" not in reasoning and "Revived" not in reasoning:
                        issues.setdefault(
                            page_id,
                            f"At {stage} but no Stage {n-1} Score (broken progression).",
                        )

    # --- Auto-fix: re-process orphans
    fixed_orphans = 0
    unmatched_orphans: list[str] = []
    for orphan_id in orphan_ids:
        try:
            result = process_single_stage1(orphan_id, config)
            decision = result.get("decision", "")
            if "unmatched" in decision:
                unmatched_orphans.append(orphan_id)
                issues[orphan_id] = (f"Unmatched submission orphan: "
                                     f"name={result.get('name')!r}. "
                                     f"No candidate record found by name or email.")
            elif "merged" in decision:
                fixed_orphans += 1
            elif "blocked" in decision:
                issues[orphan_id] = result.get("reasoning", "Submission blocked")
        except Exception as e:
            issues[orphan_id] = f"Orphan re-process raised: {type(e).__name__}: {e}"

    # --- Auto-fix: re-score stale active stages (idempotent per-stage scorers)
    rescored_stages: list[int] = []
    for n in sorted(needs_rescore.keys()):
        scorer = {2: run_stage2, 3: run_stage3, 4: run_stage4, 5: run_stage5}.get(n)
        if not scorer:
            continue
        print(f"\n[health] Stage {n} has {len(needs_rescore[n])} unscored submission(s) — rescoring")
        try:
            scorer(config)
            rescored_stages.append(n)
        except Exception as e:
            for pid in needs_rescore[n]:
                issues[pid] = f"Stage {n} re-score raised: {type(e).__name__}: {e}"

    # Re-check the rescored candidates after the run — anything still without
    # a score gets flagged.
    for n, page_ids in needs_rescore.items():
        if n not in rescored_stages:
            continue
        for pid in page_ids:
            try:
                refreshed = get_page(pid)
            except Exception:
                continue
            score = (refreshed.get("properties", {}).get(f"Stage {n} Score") or {}).get("number")
            if score is None and pid not in issues:
                submitted = ((refreshed.get("properties", {}).get(f"Stage {n} Submitted At") or {}).get("date") or {}).get("start")
                issues[pid] = (f"Stage {n} submission ({submitted}) still unscored after a "
                               f"retry — investigate the scorer or the submission shape.")

    # --- Sync the Pipeline Issue property: write new, clear resolved
    flagged_count = 0
    cleared_count = 0
    for page in all_pages:
        page_id = page["id"]
        existing = _existing_pipeline_issue(page)
        new = issues.get(page_id, "")
        if new and existing != new:
            try:
                _set_pipeline_issue(page_id, new)
                flagged_count += 1
            except Exception as e:
                print(f"  WARN: could not set Pipeline Issue on {page_id}: {e}")
        elif not new and existing:
            try:
                _set_pipeline_issue(page_id, None)
                cleared_count += 1
            except Exception as e:
                print(f"  WARN: could not clear Pipeline Issue on {page_id}: {e}")

    # Summary
    print(f"\n--- Health Check Complete ---")
    print(f"Orphans found: {len(orphan_ids)} ({fixed_orphans} merged, {len(unmatched_orphans)} unmatched)")
    print(f"Stages re-scored: {sorted(rescored_stages)}")
    print(f"Pipeline issues now flagged: {flagged_count}")
    print(f"Pipeline issues cleared: {cleared_count}")
    if issues:
        print(f"\nCurrently flagged ({len(issues)}):")
        for pid, msg in issues.items():
            print(f"  ⚠ {pid}: {msg}")
    return {
        "orphans_found": len(orphan_ids),
        "orphans_fixed": fixed_orphans,
        "orphans_unmatched": len(unmatched_orphans),
        "stages_rescored": sorted(rescored_stages),
        "flagged": len(issues),
        "cleared": cleared_count,
    }


def _has_any_stage_payload(props: dict) -> bool:
    """True if the row carries any per-stage submission file or text."""
    file_props = [
        "Notion task screenshots", "Spreadsheet task screenshots",
        "Presentation task screenshots (each slide)", "Upload your task",
        "On-chain transaction screenshot", "Lightning payment screenshot",
        "Stage 5 BTC Screenshot", "Stage 5 Lightning Screenshot",
    ]
    text_props = [
        "Stage 3 Submission", "Stage 4 Submission",
        "A.I. email draft (3/3) - edited version",
        "On-chain transaction ID", "Lightning transaction proof of payment",
        "Bitcoin — What is it and how does it work at a high level?",
    ]
    for k in file_props:
        if (props.get(k) or {}).get("files"):
            return True
    for k in text_props:
        rt = (props.get(k) or {}).get("rich_text") or []
        if "".join(t.get("plain_text", "") for t in rt).strip():
            return True
    return False


def _utcnow_for_health():
    from datetime import datetime, timezone
    return datetime.now(timezone.utc)


def _minutes_since(iso_ts: str, now_dt) -> float | None:
    from datetime import datetime
    try:
        dt = datetime.fromisoformat(iso_ts.replace("Z", "+00:00"))
        return (now_dt - dt).total_seconds() / 60
    except Exception:
        return None


# --- Final Ranking ---

def run_ranking(config: dict) -> dict:
    """Generate final ranked shortlist with normalized composite scores."""
    db_id = config["notion_database_id"]
    weights = config["composite_weights"]
    max_scores = config.get("max_scores", {"stage1": 100, "stage2": 20, "stage3": 35, "stage4": 25})
    min_finalists = config["min_finalists"]
    stages = config["stages"]

    finalists = get_candidates(db_id, stage=stages["final_interview"])
    if not finalists:
        print("No candidates in 'Final Interview' stage.")
        return {"ranked": 0}

    prompts = get_prompts(config)
    ranking_prompt = prompts[10] if len(prompts) >= 11 else None

    candidates_for_ranking = []
    for page in finalists:
        props = page["properties"]
        candidate = get_candidate_data(page)

        s1 = props.get("AI Score Stage 1", {}).get("number") or 0
        s2 = props.get("Stage 2 Score", {}).get("number") or 0
        s3 = props.get("Stage 3 Score", {}).get("number") or 0
        s4 = props.get("Stage 4 Score", {}).get("number") or 0

        # Normalize to 0-100
        s1_norm = (s1 / max_scores["stage1"]) * 100
        s2_norm = (s2 / max_scores["stage2"]) * 100
        s3_norm = (s3 / max_scores["stage3"]) * 100
        s4_norm = (s4 / max_scores["stage4"]) * 100

        # Composite (interview defaults to 0 until entered)
        composite = (
            s1_norm * weights.get("stage1", 0)
            + s2_norm * weights.get("stage2", 0)
            + s3_norm * weights.get("stage3", 0)
            + s4_norm * weights.get("stage4", 0)
        )

        candidates_for_ranking.append({
            "page_id": candidate["page_id"],
            "name": candidate["full_name"],
            "stage1_score": s1,
            "stage2_score": s2,
            "stage3_score": s3,
            "stage4_score": s4,
            "composite_score": round(composite, 1),
        })

        update_candidate(candidate["page_id"], {"Final Score": round(composite, 1)})
        time.sleep(0.35)

    candidates_for_ranking.sort(key=lambda c: c["composite_score"], reverse=True)

    if ranking_prompt:
        ai_ranking = evaluate_ranking(candidates_for_ranking, ranking_prompt)
        print("\n--- AI Ranking Result ---")
        shortlist = ai_ranking.get("shortlist", [])
        for entry in shortlist:
            print(f"  #{entry.get('rank', '?')}: {entry.get('name', '?')} — Score: {entry.get('final_score', '?')}")
            print(f"    {entry.get('summary', '')}")
        if ai_ranking.get("notes"):
            print(f"\nNotes: {ai_ranking['notes']}")
    else:
        print("\n--- Composite Score Ranking ---")
        for i, c in enumerate(candidates_for_ranking, 1):
            print(f"  #{i}: {c['name']} — Composite: {c['composite_score']}")
            print(f"    S1:{c['stage1_score']}/100  S2:{c['stage2_score']}/20  S3:{c['stage3_score']}/35  S4:{c['stage4_score']}/25")

    if len(candidates_for_ranking) < min_finalists:
        print(f"\nWARNING: Only {len(candidates_for_ranking)} finalists. Minimum is {min_finalists}.")
        print("Consider reviewing candidates in 'Manual Review' stage.")

    return {"ranked": len(candidates_for_ranking)}
