"""One-off: merge Stage 3 task file submissions into the original candidate's
Stage 3 Submission text field, then archive the orphan.

Background
----------
The Stage 3 Notion form puts the candidate's task document into the
"Upload your task" file property of a NEW row in the Candidate Applications
database (orphan). The Stage 3 scorer reads the "Stage 3 Submission" text
property — it cannot read a DOCX/PDF directly. This script extracts the file
text, writes it onto the original candidate's row, sets Stage 3 Submitted At,
and archives the orphan so run_stage3 picks the candidate up cleanly.
"""

import os
import sys
import time
from pathlib import Path

import httpx
import pypdf
import docx
from dotenv import load_dotenv

load_dotenv(override=True)

NOTION_API_KEY = os.getenv("NOTION_API_KEY")
if not NOTION_API_KEY:
    print("ERROR: NOTION_API_KEY not set")
    sys.exit(1)

NOTION_VERSION = "2022-06-28"
BASE_URL = "https://api.notion.com/v1"
HEADERS = {"Authorization": f"Bearer {NOTION_API_KEY}", "Notion-Version": NOTION_VERSION}
JSON_HEADERS = {**HEADERS, "Content-Type": "application/json"}

BACKUP_ROOT = Path(__file__).parent / "submission_backups" / "stage3"

MERGES = [
    {
        "orphan_id": "349eddaa-d74c-81cd-bfde-fc2b171ce0f6",
        "original_id": "343eddaa-d74c-81bf-9990-d4bd52436266",
        "label": "Patrick Cisuaka Beya",
    },
    {
        "orphan_id": "348eddaa-d74c-81ce-a021-f818b318ec14",
        "original_id": "343eddaa-d74c-81ec-9353-c6e35601b9e1",
        "label": "Mohammed Barday",
    },
]


def fetch_page(page_id: str) -> dict:
    r = httpx.get(f"{BASE_URL}/pages/{page_id}", headers=JSON_HEADERS, timeout=30)
    r.raise_for_status()
    return r.json()


def archive_page(page_id: str) -> None:
    r = httpx.patch(f"{BASE_URL}/pages/{page_id}", headers=JSON_HEADERS,
                    json={"archived": True}, timeout=30)
    r.raise_for_status()


def patch_props(page_id: str, properties: dict) -> None:
    r = httpx.patch(f"{BASE_URL}/pages/{page_id}", headers=JSON_HEADERS,
                    json={"properties": properties}, timeout=60)
    if r.status_code >= 400:
        print(f"  Patch failed: {r.status_code} {r.text}")
    r.raise_for_status()


def download(url: str, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with httpx.stream("GET", url, timeout=60, follow_redirects=True) as r:
        r.raise_for_status()
        with open(out_path, "wb") as f:
            for chunk in r.iter_bytes():
                f.write(chunk)


def extract_pdf_text(path: Path) -> str:
    reader = pypdf.PdfReader(str(path))
    return "\n\n".join((p.extract_text() or "").strip() for p in reader.pages).strip()


def extract_docx_text(path: Path) -> str:
    d = docx.Document(str(path))
    parts = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # Also capture text inside tables (some candidates use tables for layout)
    for table in d.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts).strip()


def chunked_rich_text(text: str, chunk_size: int = 2000) -> list[dict]:
    return [{"type": "text", "text": {"content": text[i:i + chunk_size]}}
            for i in range(0, len(text), chunk_size)]


def process(merge: dict) -> None:
    label = merge["label"]
    orphan_id = merge["orphan_id"]
    original_id = merge["original_id"]
    print(f"\n=== {label} ===")
    print(f"  Orphan:   {orphan_id}")
    print(f"  Original: {original_id}")

    orphan = fetch_page(orphan_id)
    files = (orphan["properties"].get("Upload your task") or {}).get("files") or []
    if not files:
        print("  No file in 'Upload your task'. Skipping.")
        return

    f = files[0]
    name = f.get("name", "submission")
    if f.get("type") == "file":
        url = f["file"]["url"]
    elif f.get("type") == "external":
        url = f["external"]["url"]
    else:
        print(f"  Unknown file type: {f.get('type')}. Skipping.")
        return

    backup_dir = BACKUP_ROOT / label.replace(" ", "_")
    local_path = backup_dir / name
    print(f"  Downloading {name}...")
    download(url, local_path)
    print(f"    Saved to {local_path}")

    ext = name.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        print("  Extracting PDF text...")
        text = extract_pdf_text(local_path)
    elif ext == "docx":
        print("  Extracting DOCX text...")
        text = extract_docx_text(local_path)
    else:
        print(f"  Unsupported file type: {ext}. Skipping.")
        return

    print(f"  Extracted {len(text)} chars")
    if not text:
        print("  WARNING: extraction returned empty text. Skipping (no archive).")
        return

    # Preview a snippet so the run is auditable from logs
    snippet = text[:200].replace("\n", " ")
    print(f"  Preview: {snippet}...")

    submitted_at = orphan["created_time"]

    print(f"  Patching original page (Stage 3 Submission + Submitted At)...")
    patch_props(original_id, {
        "Stage 3 Submission": {"rich_text": chunked_rich_text(text)},
        "Stage 3 Submitted At": {"date": {"start": submitted_at}},
    })

    print(f"  Archiving orphan {orphan_id}...")
    archive_page(orphan_id)
    print("  Done.")


def main() -> None:
    for m in MERGES:
        try:
            process(m)
        except httpx.HTTPStatusError as e:
            print(f"  HTTP ERROR: {e.response.status_code} {e.response.text}")
        except Exception as e:
            print(f"  ERROR: {type(e).__name__}: {e}")
        time.sleep(0.5)


if __name__ == "__main__":
    main()
